import math
import requests
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import io

class CiberPerfilado:
    def __init__(self, D1, D2, D3, D4, D5, D6, alpha1=0.2, alpha2=0.2, alpha3=0.2, alpha4=0.2, alpha5=0.1, alpha6=0.1, beta=0.5, lambda_param=0.3):
        # Asignación de las dimensiones
        self.D1 = D1
        self.D2 = D2
        self.D3 = D3
        self.D4 = D4
        self.D5 = D5
        self.D6 = D6
        
        # Asignación de los pesos de las dimensiones (coeficientes alpha)
        self.alpha1 = alpha1
        self.alpha2 = alpha2
        self.alpha3 = alpha3
        self.alpha4 = alpha4
        self.alpha5 = alpha5
        self.alpha6 = alpha6
        
        # Parámetro beta para ajustar el impacto de la centralidad de intermediación
        self.beta = beta
        
        # Parámetro lambda para el decaimiento exponencial
        self.lambda_param = lambda_param
    
    def calcular_perfil_usuario(self):
        # Cálculo del perfil de usuario P_u
        P_u = (self.alpha1 * self.D1 + 
               self.alpha2 * self.D2 + 
               self.alpha3 * self.D3 + 
               self.alpha4 * self.D4 + 
               self.alpha5 * self.D5 + 
               self.alpha6 * self.D6)
        return P_u
    
    def calcular_influencia(self, C_b):
        # Cálculo de la influencia del usuario I_u considerando la centralidad de intermediación C_b
        P_u = self.calcular_perfil_usuario()
        I_u = P_u * (1 + self.beta * C_b)
        return I_u
    
    def calcular_influencia_multigrado(self, C_b, grado):
        # Cálculo de la influencia a través de múltiples grados I_u^(n)
        I_u = self.calcular_influencia(C_b)
        I_u_n = I_u * math.exp(-self.lambda_param * grado)
        return I_u_n
    
    def calcular_impacto_total(self, C_b):
        # Cálculo del impacto total T_u sumando las influencias hasta el tercer grado
        I_u = self.calcular_influencia(C_b)
        T_u = I_u * (1 + math.exp(-self.lambda_param) + 
                     math.exp(-2 * self.lambda_param) + 
                     math.exp(-3 * self.lambda_param))
        return T_u

def obtener_datos_api():
    # Simularemos una llamada a la API con datos de ejemplo
    # En un caso real, reemplazarías esto con una llamada real a tu API
    url = "https://api.ejemplo.com/datos_usuario"
    try:
        response = requests.get(url)
        response.raise_for_status()
        datos = response.json()
        return datos
    except requests.RequestException as e:
        print(f"Error al obtener datos de la API: {e}")
        # Datos de ejemplo en caso de error
        return {
            "D1": 0.8, "D2": 0.6, "D3": 0.7, "D4": 0.5, "D5": 0.9, "D6": 0.4,
            "C_b": 0.7
        }

def crear_grafica_radial(resultados):
    # Preparar los datos para la gráfica radial
    categorias = ['D1', 'D2', 'D3', 'D4', 'D5', 'D6']
    valores = [resultados[cat] for cat in categorias]

    # Crear la figura y los ejes polares
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(projection='polar'))

    # Número de categorías
    N = len(categorias)

    # Calcular los ángulos para cada categoría
    angulos = [n / float(N) * 2 * np.pi for n in range(N)]
    valores += valores[:1]  # Repetir el primer valor para cerrar el polígono
    angulos += angulos[:1]

    # Dibujar la gráfica
    ax.plot(angulos, valores)
    ax.fill(angulos, valores, alpha=0.3)

    # Establecer las etiquetas y los límites
    ax.set_xticks(angulos[:-1])
    ax.set_xticklabels(categorias)
    ax.set_ylim(0, 1)

    # Añadir título
    plt.title('Gráfica Radial de Dimensiones')

    # Guardar la gráfica en un buffer de bytes
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    
    plt.close()  # Cerrar la figura para liberar memoria
    return buf

def exportar_a_docx(resultados):
    doc = Document()
    doc.add_heading('Resultados del Análisis de CiberPerfilado', 0)

    doc.add_heading('Datos de entrada', level=1)
    p = doc.add_paragraph()
    p.add_run('Dimensiones: ').bold = True
    p.add_run(f"D1 = {resultados['D1']}, D2 = {resultados['D2']}, D3 = {resultados['D3']}, "
              f"D4 = {resultados['D4']}, D5 = {resultados['D5']}, D6 = {resultados['D6']}")
    
    p = doc.add_paragraph()
    p.add_run('Pesos (alpha): ').bold = True
    p.add_run(f"α1 = {resultados['alpha1']}, α2 = {resultados['alpha2']}, α3 = {resultados['alpha3']}, "
              f"α4 = {resultados['alpha4']}, α5 = {resultados['alpha5']}, α6 = {resultados['alpha6']}")
    
    p = doc.add_paragraph()
    p.add_run('Beta (β): ').bold = True
    p.add_run(f"{resultados['beta']}")
    
    p = doc.add_paragraph()
    p.add_run('Lambda (λ): ').bold = True
    p.add_run(f"{resultados['lambda_param']}")
    
    p = doc.add_paragraph()
    p.add_run('Centralidad de intermediación (C_b): ').bold = True
    p.add_run(f"{resultados['C_b']}")

    doc.add_heading('Resultados calculados', level=1)
    for key, value in resultados['calculos'].items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(f"{value:.4f}")

    doc.add_heading('Interpretación', level=1)
    doc.add_paragraph('El usuario tiene un perfil base sólido.')
    doc.add_paragraph('Su influencia directa es significativamente mayor debido a su alta centralidad de intermediación.')
    doc.add_paragraph('La influencia disminuye con la distancia en la red, como se ve en la influencia de segundo grado.')
    doc.add_paragraph('El impacto total sugiere que el usuario tiene una influencia considerable que se extiende más allá de sus conexiones inmediatas.')

    # Añadir la gráfica radial
    doc.add_heading('Gráfica Radial de Dimensiones', level=1)
    grafica_buf = crear_grafica_radial(resultados)
    doc.add_picture(grafica_buf, width=Inches(6))

    doc.save('Resultados_CiberPerfilado.docx')
    print("Documento 'Resultados_CiberPerfilado.docx' creado exitosamente.")

# Obtener datos de la API
datos_api = obtener_datos_api()

# Crear una instancia de CiberPerfilado con los datos de la API
perfil = CiberPerfilado(
    D1=datos_api['D1'], D2=datos_api['D2'], D3=datos_api['D3'],
    D4=datos_api['D4'], D5=datos_api['D5'], D6=datos_api['D6'],
    alpha1=0.25, alpha2=0.2, alpha3=0.15, alpha4=0.15, alpha5=0.15, alpha6=0.1,
    beta=0.6, lambda_param=0.25
)

# Calcular los diferentes valores
C_b = datos_api['C_b']
perfil_usuario = perfil.calcular_perfil_usuario()
influencia = perfil.calcular_influencia(C_b)
influencia_grado2 = perfil.calcular_influencia_multigrado(C_b, 2)
impacto_total = perfil.calcular_impacto_total(C_b)

# Preparar los resultados para exportar
resultados = {
    'D1': datos_api['D1'], 'D2': datos_api['D2'], 'D3': datos_api['D3'],
    'D4': datos_api['D4'], 'D5': datos_api['D5'], 'D6': datos_api['D6'],
    'alpha1': 0.25, 'alpha2': 0.2, 'alpha3': 0.15, 'alpha4': 0.15, 'alpha5': 0.15, 'alpha6': 0.1,
    'beta': 0.6, 'lambda_param': 0.25, 'C_b': C_b,
    'calculos': {
        'Perfil de usuario (P_u)': perfil_usuario,
        'Influencia del usuario (I_u)': influencia,
        'Influencia de segundo grado (I_u^(2))': influencia_grado2,
        'Impacto total (T_u)': impacto_total
    }
}

# Exportar resultados a un documento Word
exportar_a_docx(resultados)

# Imprimir los resultados en la consola
print(f"Perfil de usuario (P_u): {perfil_usuario:.4f}")
print(f"Influencia del usuario (I_u): {influencia:.4f}")
print(f"Influencia de segundo grado (I_u^(2)): {influencia_grado2:.4f}")
print(f"Impacto total (T_u): {impacto_total:.4f}")
